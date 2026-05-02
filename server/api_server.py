from flask import Flask, request, jsonify
import os
import sys
import requests
from google import genai
import json
from flask_cors import CORS
from typing import Dict, Any
import logging
import hashlib


sys.path.append(os.path.dirname(__file__))
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'Agents'))
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))  # For logging_config

# Original imports
from resume_tailoring_agent import get_google_services, get_doc_id_from_url, read_google_doc_content
from job_application_agent import run_links_with_refactored_agent
from logging_config import setup_file_logging
from google_oauth_service import GoogleOAuthService

# Production infrastructure imports
from rate_limiter import rate_limiter
from job_queue import job_queue
from security_manager import require_secure_headers
from database_optimizer import setup_database_optimizations
from backup_manager import schedule_backups
from bug_bounty import (
    get_user_bonus_for_limit,
    effective_limit,
)
from routes.account import create_account_blueprint
from routes.auth import create_auth_blueprint
from routes.beta import create_beta_blueprint
from routes.cli import create_cli_blueprint
from routes.feedback import create_feedback_blueprint
from routes.health import create_health_blueprint
from routes.jobs import create_jobs_blueprint
from routes.monitoring import create_monitoring_blueprint
from routes.oauth import create_frontend_redirect_builder, create_oauth_blueprint
from routes.profile import create_profile_blueprint
from routes.projects import create_projects_blueprint
from routes.public import create_public_blueprint
from routes.resume import create_resume_blueprint
from routes.settings import create_settings_blueprint


#Initialize the app
app = Flask(__name__)

# ============= RESOURCE MANAGEMENT & MONITORING SETUP =============
# Initialize resource manager, connection pool, and health monitor
try:
    from health_monitor import initialize_system

    # Initialize all resource management components
    initialize_system()
    logging.info("✅ Resource management and monitoring initialized")
    
except ImportError as e:
    logging.warning(f"⚠️ Resource management not available: {e}")
    logging.info("   System will run without advanced resource management")
except Exception as e:
    logging.error(f"❌ Failed to initialize resource management: {e}")
    import traceback
    logging.error(traceback.format_exc())

# ============= END RESOURCE MANAGEMENT SETUP =============

# ============= SENTRY ERROR TRACKING =============
# Initialize Sentry for production error tracking
SENTRY_ENABLED = False
try:
    import sentry_sdk
    from sentry_sdk.integrations.flask import FlaskIntegration

    sentry_dsn = os.getenv('SENTRY_DSN')
    sentry_environment = os.getenv('FLASK_ENV', 'development')

    if sentry_dsn:
        sentry_sdk.init(
            dsn=sentry_dsn,
            integrations=[FlaskIntegration()],
            environment=sentry_environment,
            traces_sample_rate=0.1,  # 10% of transactions for performance monitoring
            profiles_sample_rate=0.1,  # 10% for profiling
            # Filter out health check endpoints from traces
            before_send_transaction=lambda event, hint: None if event.get('transaction', '').startswith('/health') or event.get('transaction', '').startswith('/ready') else event,
        )
        SENTRY_ENABLED = True
        logging.info(f"✅ Sentry error tracking initialized (environment: {sentry_environment})")
    else:
        logging.info("⚠️ Sentry DSN not configured - error tracking disabled")

except ImportError:
    logging.info("⚠️ Sentry SDK not installed - error tracking disabled")
    logging.info("   Install with: pip install sentry-sdk[flask]")

# ============= END SENTRY SETUP =============

# Register modular route blueprints (incremental extraction from monolithic API).
app.register_blueprint(create_health_blueprint(sentry_enabled=SENTRY_ENABLED))
app.register_blueprint(create_profile_blueprint())
app.register_blueprint(create_beta_blueprint())
app.register_blueprint(create_account_blueprint())
app.register_blueprint(
    create_feedback_blueprint(invalidate_credits_cache=lambda user_id: _invalidate_credits_cache(user_id))
)
app.register_blueprint(create_projects_blueprint())
app.register_blueprint(create_monitoring_blueprint())
app.register_blueprint(create_cli_blueprint())
app.register_blueprint(create_settings_blueprint())
app.register_blueprint(
    create_resume_blueprint(
        extract_google_doc_with_oauth=lambda resume_url, user_id: extract_google_doc_with_oauth(resume_url, user_id),
        process_resume_with_llm=lambda resume_text: process_resume_with_llm(resume_text),
        compute_resume_text_hash=lambda resume_text: _compute_resume_text_hash(resume_text),
    )
)
app.register_blueprint(
    create_jobs_blueprint(
        get_user_and_limit=lambda user_id, limit_type: _get_user_and_limit(user_id, limit_type),
        invalidate_credits_cache=lambda user_id: _invalidate_credits_cache(user_id),
    )
)
app.register_blueprint(create_public_blueprint())

# Configure CORS for development and production
# Default includes multiple localhost ports for development and Vercel production
default_origins = 'http://localhost:3000,http://localhost:3001,http://localhost:5173,https://job-agent-frontend-two.vercel.app'
allowed_origins_str = os.getenv('CORS_ORIGINS', default_origins)

# Parse allowed origins
allowed_origins = [origin.strip() for origin in allowed_origins_str.split(',')]

flask_env = os.getenv('FLASK_ENV')
logging.info(f"✅ CORS: Explicit origins configured: {len(allowed_origins)} origins")

# Apply CORS with expanded origins list (supports regex for Vercel in dev only)
CORS(
    app, 
    origins=allowed_origins, 
    supports_credentials=True,
    allow_headers=['Content-Type', 'Authorization', 'Accept'],
    methods=['GET', 'POST', 'PUT', 'DELETE', 'OPTIONS', 'PATCH'],
    expose_headers=['Content-Type', 'Authorization']
)

_frontend_redirect_builder = create_frontend_redirect_builder(
    allowed_origins=allowed_origins,
    flask_env=flask_env,
)
app.register_blueprint(create_auth_blueprint(build_frontend_redirect=_frontend_redirect_builder))
app.register_blueprint(
    create_oauth_blueprint(
        allowed_origins=allowed_origins,
        flask_env=flask_env,
    )
)

# Global handler for CORS preflight OPTIONS requests
@app.before_request
def handle_preflight():
    if request.method == 'OPTIONS':
        response = app.make_default_options_response()
        return response

# Apply security headers to all responses
@app.after_request
@require_secure_headers
def after_request(response):
    # The /pick-resume page embeds inline styles, inline scripts, and loads
    # external resources from Google (apis.google.com, accounts.google.com,
    # content.googleapis.com).  The blanket "default-src 'self'" CSP blocks
    # all of that, so we replace it with a permissive-but-scoped policy for
    # that one route only.
    if request.path == "/pick-resume":
        response.headers['Content-Security-Policy'] = (
            "default-src 'self'; "
            "style-src 'self' 'unsafe-inline'; "
            "script-src 'self' 'unsafe-inline' https://apis.google.com https://accounts.google.com; "
            "frame-src https://docs.google.com https://drive.google.com https://accounts.google.com; "
            "connect-src 'self' https://accounts.google.com https://www.googleapis.com; "
            "img-src 'self' data: https://ssl.gstatic.com https://www.gstatic.com;"
        )
    return response

JOBS: Dict[str, Dict[str, Any]] = {}

# Initialize production infrastructure
def initialize_production_infrastructure():
    """Initialize all production infrastructure components"""
    try:
        # Initialize database tables if they don't exist
        from database_config import Base, engine, test_connection, _apply_incremental_migrations
        logging.info("Checking database connection...")
        if test_connection():
            logging.info("✅ Database connection successful")
            logging.info("Initializing database tables...")
            Base.metadata.create_all(bind=engine)
            _apply_incremental_migrations()
            logging.info("✅ Database tables initialized")
        else:
            raise Exception("Database connection failed")
        
        # Set up database optimizations
        setup_database_optimizations()
        logging.info("✅ Database optimizations initialized")
        
        # Start job queue worker
        job_queue.start_worker()
        logging.info("✅ Job queue worker started")
        
        # Schedule automated backups
        schedule_backups()
        logging.info("✅ Backup scheduler initialized")
        
        logging.info("🚀 Production infrastructure initialized successfully")
        
    except Exception as e:
        logging.error(f"❌ Failed to initialize production infrastructure: {e}")
        raise

def get_shared_gemini_api_key() -> str:
    return (
        os.getenv("LAUNCHWAY_GEMINI_API_KEY")
        or os.getenv("GEMINI_API_KEY")
        or os.getenv("GOOGLE_API_KEY")
        or ""
    )


def initialize_gemini():
    api_key = get_shared_gemini_api_key()
    return genai.Client(api_key=api_key)

def process_resume_with_llm(resume_text: str) -> Dict[str, Any]:
    from Agents.gemini_key_manager import GeminiKeyManager, GeminiQuotaExhaustedError
    _key_mgr = GeminiKeyManager(
        primary_mode="launchway",
        secondary_mode=None,
        launchway_api_key=get_shared_gemini_api_key(),
        cooldown_seconds=15,
    )
    profile_schema={
            "type": "object",
            "properties": {
                "first name": {"type": "string"},
                "last name": {"type": "string"},
                "email": {"type": "string"},
                "date of birth": {"type": "string"},
                "phone": {"type": "string"},
                "address": {"type": "string"},
                "city": {"type": "string"},
                "state": {"type": "string"},
                "zip": {"type": "string"},
                "country": {"type": "string"},
                "country_code": {"type": "string"},
                "state_code": {"type": "string"},
                "linkedin": {"type": "string"},
                "github": {"type": "string"},
                "other links": {
                    "type": "array",
                    "items": {"type": "string"}
                },
                "education": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "degree": {"type": "string"},
                            "institution": {"type": "string"},
                            "graduation_year": {"type": "string"},
                            "gpa": {"type": "string"},
                            "relevant_courses": {
                                "type": "array",
                                "items": {"type": "string"}
                            }
                        }
                    }
                },
                "work experience": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string"},
                            "company": {"type": "string"},
                            "start_date": {"type": "string"},
                            "end_date": {"type": "string"},
                            "description": {"type": "string"},
                            "achievements": {
                                "type": "array",
                                "items": {"type": "string"}
                            }
                        }
                    }
                },
                "projects": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "description": {"type": "string"},
                            "technologies": {
                                "type": "array",
                                "items": {"type": "string"}
                            },
                            "github_url": {"type": "string"},
                            "live_url": {"type": "string"},
                            "features": {
                                "type": "array",
                                "items": {"type": "string"}
                            }
                        }
                    }
                },
                "skills": {
                    "type": "object",
                    "properties": {
                        "technical": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "programming_languages": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "frameworks": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "tools": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "soft_skills": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "languages": {
                            "type": "array",
                            "items": {"type": "string"}
                        }
                    }
                },
                "summary": {"type": "string"}
            }
        }
    # Create a template JSON structure for Gemini to fill
    json_template = {
        "first name": "",
        "last name": "",
        "email": "",
        "phone": "",
        "address": "",
        "city": "",
        "state": "",
        "zip": "",
        "country": "",
        "country_code": "",
        "state_code": "",
        "linkedin": "",
        "github": "",
        "other links": [],
        "date of birth": "",
        "education": [
            {
                "degree": "",
                "institution": "",
                "graduation_year": "",
                "gpa": "",
                "relevant_courses": []
            }
        ],
        "work experience": [
            {
                "title": "",
                "company": "",
                "start_date": "",
                "end_date": "",
                "description": "",
                "achievements": []
            }
        ],
        "projects": [
            {
                "name": "",
                "description": "",
                "technologies": [],
                "github_url": "",
                "live_url": "",
                "features": []
            }
        ],
        "skills": {
            "technical": [],
            "programming_languages": [],
            "frameworks": [],
            "tools": [],
            "soft_skills": [],
            "languages": []
        },
        "summary": ""
    }

    prompt = f"""
    You are an expert at extracting information from resumes and structuring it into a standardized profile format.

    Please analyze the following resume text and extract the relevant information to populate the JSON template below.

    Resume Text:
    {resume_text}

    Fill in the following JSON template with the extracted information. Replace empty strings and empty arrays with actual data from the resume:

    {json.dumps(json_template, indent=2)}

    Instructions:
    1. Extract the person's name and split it into first name and last name
    2. Look for contact information (email, phone, address, city, state, zip, country)
    3. Extract education history - create multiple entries in the education array if there are multiple degrees
    4. Extract work experience - create multiple entries in the work experience array for each job
    5. Extract projects - create multiple entries in the projects array for each project
    6. Extract and categorize skills:
       - technical: Core technical skills (e.g., "Machine Learning", "Data Science", "Cloud Computing")
       - programming_languages: Programming languages (e.g., "Python", "JavaScript", "Java", "C++")
       - frameworks: Frameworks and libraries (e.g., "React", "Django", "TensorFlow", "Spring")
       - tools: Tools and technologies (e.g., "AWS", "Docker", "Git", "PostgreSQL", "MongoDB")
       - soft_skills: Soft skills (e.g., "Leadership", "Communication", "Problem Solving")
       - languages: Spoken languages with proficiency (e.g., "English (Native)", "Spanish (Fluent)")
    7. Create a professional summary (2-3 sentences highlighting key strengths)
    8. Look for LinkedIn, GitHub, and other professional links
    9. Use full official names for countries and states (e.g., "United States of America", "California")
    10. Leave fields empty if information is not available

    Return ONLY the filled JSON object (not an array). Do not include any markdown formatting, code blocks, or additional text.
    """

    logging.info("Sending resume text to Gemini for profile extraction")

    try:
        response = _key_mgr.generate_content(
            "gemini-2.5-flash",
            prompt,
            config={
                "response_mime_type": "application/json",
                "thinking_config": {"thinking_budget": 0},
                "max_output_tokens": 8192,
            },
        )

        # Clean the response text to extract JSON
        response_text = response.text.strip()

        # Remove any markdown formatting if present
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        if response_text.startswith('```'):
            response_text = response_text[3:]

        response_text = response_text.strip()

        try:
            profile_data = json.loads(response_text)

            # Handle case where Gemini returns an array instead of object
            if isinstance(profile_data, list) and len(profile_data) > 0:
                profile_data = profile_data[0]
            elif not isinstance(profile_data, dict):
                logging.error(f"process_resume_with_llm: unexpected response type {type(profile_data)}")
                return None

            validated_data = _validate_profile_data(profile_data, profile_schema)
            logging.info("Resume profile extraction successful")
            return validated_data
        except json.JSONDecodeError as json_err:
            logging.error(f"process_resume_with_llm: JSON parse error: {json_err}")
            logging.debug(f"process_resume_with_llm: raw response text (first 500 chars): {response_text[:500]!r}")
            return None
    except GeminiQuotaExhaustedError as qe:
        logging.error(f"process_resume_with_llm: Gemini quota exhausted after retries: {qe}")
        return None
    except Exception as e:
        logging.error(f"process_resume_with_llm: Gemini error: {e}")
        return None

def _validate_profile_data(profile_data: Dict[str, Any], schema: Dict[str, Any]) -> Dict[str, Any]:
    """Validate and clean the profile data to match the schema"""
    # Extract the actual properties from the JSON schema
    if "properties" in schema:
        schema_properties = schema["properties"]
    else:
        schema_properties = schema
    
    validated_data = {}
    
    # Define default values for each field type
    default_values = {
        "first name": "",
        "last name": "",
        "email": "",
        "date of birth": "",
        "phone": "",
        "address": "",
        "city": "",
        "state": "",
        "zip": "",
        "country": "",
        "country_code": "",
        "state_code": "",
        "linkedin": "",
        "github": "",
        "other links": [],
        "education": [],
        "work experience": [],
        "projects": [],
        "skills": {
            "technical": [],
            "programming_languages": [],
            "frameworks": [],
            "tools": [],
            "soft_skills": [],
            "languages": []
        },
        "summary": ""
    }
    
    for key, default_value in default_values.items():
        if key in profile_data:
            value = profile_data[key]
            
            # Handle None values
            if value is None:
                validated_data[key] = default_value
            # Validate based on expected type
            elif isinstance(default_value, str):
                validated_data[key] = str(value) if value else ""
            elif isinstance(default_value, list):
                if isinstance(value, list):
                    validated_data[key] = value
                else:
                    validated_data[key] = []
            elif isinstance(default_value, dict):
                if isinstance(value, dict):
                    validated_data[key] = value
                else:
                    validated_data[key] = default_value
            else:
                validated_data[key] = value
        else:
            validated_data[key] = default_value
    
    return validated_data

def _normalize_resume_text_for_hash(resume_text: str) -> str:
    """Normalize resume text to detect meaningful content changes."""
    if not resume_text:
        return ""
    return " ".join(resume_text.split()).strip().lower()

def _compute_resume_text_hash(resume_text: str) -> str:
    """Compute a stable hash for resume keyword cache invalidation."""
    normalized = _normalize_resume_text_for_hash(resume_text)
    if not normalized:
        return ""
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()

def extract_google_doc_with_oauth(resume_url: str, user_id: int) -> str:
    """
    Extract text from Google Doc using OAuth credentials (works with private docs)
    Falls back to public access if OAuth not connected
    """
    try:
        # Check if user has Google OAuth connected
        if GoogleOAuthService.is_connected(user_id):
            # Prefer OAuth access for connected users (supports private docs).
            try:
                logging.info(f"User {user_id} has OAuth connected, using authenticated access")
                credentials = GoogleOAuthService.get_credentials(user_id)

                doc_id = get_doc_id_from_url(resume_url)
                docs_service, _ = get_google_services(credentials)
                resume_text = read_google_doc_content(docs_service, doc_id)

                if resume_text and resume_text.strip():
                    logging.info(f"Successfully read private Google Doc via OAuth for user {user_id}")
                    return resume_text

                # OAuth returned empty text after retries - transient Google API issue.
                logging.warning(f"OAuth read returned empty text for user {user_id}")
                raise ValueError(
                    "Google's API returned an empty response. This is usually temporary - "
                    "please wait a moment and try again."
                )
            except ValueError:
                raise  # propagate our own clear messages unchanged
            except Exception as oauth_err:
                logging.warning(f"OAuth access failed for user {user_id}: {oauth_err}")
                oauth_err_text = str(oauth_err).lower()

                # Recover from revoked/expired Google tokens. This typically appears as
                # invalid_grant during refresh. In that case, clear stored OAuth tokens
                # and try public-link extraction as a graceful fallback.
                if "invalid_grant" in oauth_err_text:
                    try:
                        GoogleOAuthService.disconnect_google_account(user_id)
                        logging.info(f"Cleared stale Google OAuth tokens for user {user_id}")
                    except Exception as disconnect_err:
                        logging.warning(
                            f"Failed to clear stale Google OAuth tokens for user {user_id}: {disconnect_err}"
                        )

                    try:
                        logging.info(f"Falling back to public Google Doc access for user {user_id}")
                        return extract_resume_text(resume_url)
                    except Exception as public_err:
                        raise ValueError(
                            "Your Google account connection expired. Reconnect Google in the Resume menu, "
                            f"or make the doc viewable by link. Public fallback failed: {public_err}"
                        )

                raise ValueError(
                    f"Could not read your Google Doc (temporary error: {oauth_err}). "
                    "Please try again in a moment."
                )

        # No OAuth - try public access (doc must be shared as 'Anyone with the link')
        return extract_resume_text(resume_url)

    except Exception as e:
        logging.error(f"Error extracting Google Doc: {e}")
        raise

def extract_resume_text(resume_url: str) -> str:
    """Extract text from publicly accessible Google Doc (legacy method)"""
    try:
        print(f"Starting to extract text from: {resume_url}")
        from urllib.parse import urlparse, parse_qs
        #Extract the document ID from the URL
        parsed_url = urlparse(resume_url)
        print(f"Parsed URL: {parsed_url}")

        if 'docs.google.com' not in parsed_url.netloc:
            print(f"Invalid Google Docs URL: {resume_url}")
            raise ValueError("Invalid Google Docs URL")

       # Extract document ID from path
        path_parts = parsed_url.path.split('/')
        print(f"Path parts: {path_parts}")

        if 'document' not in path_parts or 'd' not in path_parts:
            print(f"Invalid Google Docs URL format: {resume_url}")
            raise ValueError("Invalid Google Docs URL format")

        doc_id = None
        for i, part in enumerate(path_parts):
            if part == 'd' and i + 1 < len(path_parts):
                doc_id = path_parts[i + 1]
                break

        print(f"Document ID: {doc_id}")


        if not doc_id:
            print(f"Could not extract document ID from URL: {resume_url}")
            raise ValueError("Could not extract document ID from URL")

        # Convert to export URL for plain text
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
        print(f"Export URL: {export_url}")

        # Make request to get the document content
        response = requests.get(export_url, timeout=30)
        print(f"Response: {response}")

        # Handle specific error cases
        if response.status_code in (401, 403):
            raise ValueError(
                "This Google Doc is private. Please either:\n"
                "  • Connect your Google account (option in the Resume menu), or\n"
                "  • Set sharing to 'Anyone with the link can view' in Google Docs."
            )
        elif response.status_code == 404:
            raise ValueError("Google Doc not found. Please check the URL.")

        response.raise_for_status()

        # Return the text content
        return response.text.strip()
    except ValueError:
        raise  # propagate our own clear messages unchanged
    except Exception as e:
        logging.error(f"Error fetching Google Doc: {e}")
        err = str(e)
        if "401" in err or "403" in err:
            raise ValueError(
                "This Google Doc is private. Please either connect your Google account "
                "or set sharing to 'Anyone with the link can view'."
            )
        elif "404" in err:
            raise ValueError("Google Doc not found. Please check the URL.")
        else:
            raise ValueError(f"Could not access Google Doc: {err}")

def extract_pdf_text(file_obj) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        import PyPDF2
        import io

        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_obj.read()))

        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            except Exception as page_err:
                logging.warning(f"Could not extract text from page {page_num}: {page_err}")
                continue

        text = text.strip()

        if not text:
            raise ValueError("No text could be extracted from the PDF. Please ensure the PDF contains selectable text (not scanned images).")

        logging.info(f"Successfully extracted {len(text)} characters from PDF")
        return text

    except Exception as e:
        logging.error(f"Error extracting PDF text: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_docx_text(file_obj) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        from docx import Document
        import io

        # Create a Document object from the file
        doc = Document(io.BytesIO(file_obj.read()))

        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

        text = text.strip()

        if not text:
            raise ValueError("No text could be extracted from the DOCX file. Please ensure the file is not empty.")

        logging.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text

    except ImportError:
        raise ValueError("DOCX file support not available. Please upload a PDF or Google Docs link instead.")
    except Exception as e:
        logging.error(f"Error extracting DOCX text: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

def _invalidate_credits_cache(user_id: str) -> None:
    try:
        from rate_limiter import redis_client as _rc
        _rc.delete(f"credits_cache:{user_id}")
    except Exception:
        pass


def _get_user_and_limit(user_id: str, limit_type: str):
    """Return (user_obj_or_none, effective_limit_int)."""
    from database_config import SessionLocal, User
    import uuid as uuid_module

    base_limit = int(rate_limiter.LIMITS[limit_type].requests)
    db = SessionLocal()
    try:
        user_uuid = uuid_module.UUID(str(user_id))
        user = db.query(User).filter(User.id == user_uuid).first()
        if not user:
            return None, base_limit
        bonus_limit = get_user_bonus_for_limit(user, limit_type)
        return user, effective_limit(base_limit, bonus_limit)
    except Exception:
        return None, base_limit
    finally:
        db.close()


# Auto job apply endpoint removed - feature only available via CLI

# Auto batch job apply endpoint removed - feature only available via CLI


if __name__ == "__main__":
    # Set up file logging for API server with DEBUG level to capture everything
    log_file = setup_file_logging(log_level=logging.DEBUG, console_logging=True)
    logging.info(f"API Server starting. Logs will be saved to: {log_file}")

    # Initialize production infrastructure (non-blocking - allow server to start)
    try:
        initialize_production_infrastructure()
    except Exception as e:
        logging.error(f"⚠️ Failed to initialize production infrastructure: {e}")
        logging.warning("⚠️ Server will start anyway - some features may be limited")
        logging.info("   Health check endpoint will be available")
        logging.info("   Readiness check will show component status")
        # Don't exit - let the server start so health checks work
        # Railway can detect the issue via /ready endpoint

    # ============= GRACEFUL SHUTDOWN HANDLERS =============
    import signal

    def graceful_shutdown(signum, frame):
        """Handle graceful shutdown on SIGTERM or SIGINT"""
        signal_name = 'SIGTERM' if signum == signal.SIGTERM else 'SIGINT'
        logging.info(f"🛑 Received {signal_name}, initiating graceful shutdown...")

        try:
            # Stop job queue worker
            logging.info("Stopping job queue worker...")
            job_queue.stop_worker()
            logging.info("✓ Job queue worker stopped")
        except Exception as e:
            logging.error(f"Error stopping job queue: {e}")

        logging.info("✅ Graceful shutdown complete")
        sys.exit(0)

    # Register signal handlers for graceful shutdown
    signal.signal(signal.SIGTERM, graceful_shutdown)
    signal.signal(signal.SIGINT, graceful_shutdown)
    logging.info("✅ Graceful shutdown handlers registered (SIGTERM, SIGINT)")

    # ============= END GRACEFUL SHUTDOWN HANDLERS =============

    # Check if we're in development or production mode
    import os
    is_development = os.getenv('FLASK_ENV') == 'development'

    # Get port from environment variable (Railway, Heroku, etc.) or default to 5000
    port = int(os.getenv('PORT', 5000))

    logging.info(f"🚀 Starting server on port {port}")
    logging.info(f"   Mode: {'DEVELOPMENT' if is_development else 'PRODUCTION'}")
    
    # Use standard Flask server
    app.run(
        host='0.0.0.0',
        port=port,
        debug=is_development,
        use_reloader=False
    )

    # Cleanup on shutdown
    try:
        job_queue.stop_worker()
        logging.info("Job queue worker stopped")
    except:
        pass
    
    print("API Server stopped")